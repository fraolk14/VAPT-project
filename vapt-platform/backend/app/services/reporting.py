import csv
import io
import json
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable
from urllib.parse import urlparse
from app.models.finding import Finding

try:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
except Exception:  # pragma: no cover
    colors = None
    TA_LEFT = 0
    letter = None
    getSampleStyleSheet = None
    ParagraphStyle = None
    inch = 72
    Image = object
    PageBreak = object
    Paragraph = object
    SimpleDocTemplate = object
    Spacer = object
    Table = object
    TableStyle = object

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None


SEVERITY_ORDER = {"critical": 4, "high": 3, "medium": 2, "low": 1, "info": 0}
REPORT_DATA_DIR = Path(__file__).resolve().parents[1] / "data"
REPORT_BRANDING_PATH = REPORT_DATA_DIR / "report_branding.json"
REPORT_UPLOADS_DIR = REPORT_DATA_DIR / "uploads"


def ensure_report_storage() -> None:
    REPORT_UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
    if not REPORT_BRANDING_PATH.exists():
        REPORT_BRANDING_PATH.write_text(
            json.dumps(
                {
                    "company_name": "VAP",
                    "logo_path": None,
                    "logo_name": None,
                    "updated_at": None,
                },
                indent=2,
            ),
            encoding="utf-8",
        )


def load_report_branding() -> dict:
    ensure_report_storage()
    try:
        payload = json.loads(REPORT_BRANDING_PATH.read_text(encoding="utf-8"))
    except Exception:
        payload = {}
    return {
        "company_name": payload.get("company_name") or "VAP",
        "logo_path": payload.get("logo_path"),
        "logo_name": payload.get("logo_name"),
        "updated_at": payload.get("updated_at"),
    }


def save_report_branding(*, company_name: str | None = None, logo_path: str | None = None, logo_name: str | None = None) -> dict:
    current = load_report_branding()
    if company_name:
        current["company_name"] = company_name
    if logo_path is not None:
        current["logo_path"] = logo_path
    if logo_name is not None:
        current["logo_name"] = logo_name
    current["updated_at"] = datetime.now(timezone.utc).isoformat()
    REPORT_BRANDING_PATH.write_text(json.dumps(current, indent=2), encoding="utf-8")
    return current


def _finding_target(finding: Finding) -> str:
    metadata = finding.finding_metadata or {}
    for key in ("target", "affected_url", "host", "hostname", "ip_address"):
        value = metadata.get(key)
        if value:
            return str(value)
    url = metadata.get("url")
    if url:
        hostname = urlparse(str(url)).hostname
        if hostname:
            return hostname
        return str(url)
    if finding.service:
        return f"{finding.service}:{finding.port}"
    return f"port {finding.port}"


def _serialize_finding(finding: Finding) -> dict:
    metadata = finding.finding_metadata or {}
    target = _finding_target(finding)
    return {
        "id": str(finding.id),
        "title": finding.title,
        "source": finding.source,
        "severity": finding.severity,
        "status": finding.status,
        "cvss_score": finding.cvss_score,
        "cve_id": finding.cve_id,
        "compliance_map": finding.compliance_map or [],
        "remediation": finding.remediation,
        "target": target,
        "asset_name": metadata.get("asset_name") or metadata.get("hostname") or target,
        "hostname": metadata.get("hostname") or metadata.get("host"),
        "ip_address": metadata.get("ip_address") or metadata.get("host"),
        "url": metadata.get("url") or metadata.get("affected_url"),
        "evidence": finding.evidence,
        "references": metadata.get("references", []),
        "os_family": metadata.get("os_family"),
        "cis_benchmark": metadata.get("cis_benchmark"),
        "hardening_recommendation": metadata.get("hardening_recommendation"),
        "assigned_to": finding.assigned_to,
        "team_name": finding.team_name,
        "verification_state": finding.verification_state,
        "detected_at": finding.detected_at.isoformat() if finding.detected_at else None,
        "resolved_at": finding.resolved_at.isoformat() if finding.resolved_at else None,
        "port": finding.port,
        "protocol": finding.protocol,
        "service": finding.service,
    }


def _format_dt(value: str | None) -> str:
    if not value:
        return "n/a"
    try:
        return datetime.fromisoformat(value.replace("Z", "+00:00")).strftime("%d-%m-%Y %H:%M UTC")
    except Exception:
        return value


def export_findings_json(findings: Iterable[Finding]) -> str:
    return json.dumps([_serialize_finding(finding) for finding in findings], indent=2)


def export_findings_csv(findings: Iterable[Finding]) -> str:
    buffer = io.StringIO()
    writer = csv.writer(buffer)
    writer.writerow(
        [
            "id",
            "title",
            "target",
            "asset_name",
            "hostname",
            "ip_address",
            "url",
            "source",
            "severity",
            "status",
            "verification_state",
            "cvss_score",
            "cve_id",
            "os_family",
            "cis_benchmark",
            "assigned_to",
            "team_name",
            "detected_at",
            "remediation",
            "evidence",
        ]
    )
    for finding in findings:
        row = _serialize_finding(finding)
        writer.writerow(
            [
                row["id"],
                row["title"],
                row["target"],
                row["asset_name"],
                row["hostname"],
                row["ip_address"],
                row["url"],
                row["source"],
                row["severity"],
                row["status"],
                row["verification_state"],
                row["cvss_score"],
                row["cve_id"],
                row["os_family"],
                row["cis_benchmark"],
                row["assigned_to"],
                row["team_name"],
                row["detected_at"],
                row["remediation"],
                row["evidence"],
            ]
        )
    return buffer.getvalue()


def summarize_findings(findings: Iterable[Finding]) -> dict:
    finding_list = list(findings)
    severity_counts = Counter((finding.severity or "info").lower() for finding in finding_list)
    source_counts = Counter(finding.source for finding in finding_list)
    compliance_counts = Counter(item for finding in finding_list for item in (finding.compliance_map or []))
    top_items = sorted(
        finding_list,
        key=lambda finding: ((finding.cvss_score or 0), SEVERITY_ORDER.get((finding.severity or "info").lower(), 0)),
        reverse=True,
    )[:10]
    return {
        "total_findings": len(finding_list),
        "open_findings": sum(1 for finding in finding_list if finding.status == "open"),
        "severity_counts": dict(severity_counts),
        "source_counts": dict(source_counts),
        "compliance_counts": dict(compliance_counts),
        "top_findings": [_serialize_finding(finding) for finding in top_items],
    }


def _severity_rank(severity: str | None) -> int:
    return SEVERITY_ORDER.get((severity or "info").lower(), 0)


def _scope_targets(items: list[dict]) -> list[str]:
    ordered = []
    seen = set()
    for item in items:
        target = item.get("target")
        if target and target not in seen:
            ordered.append(target)
            seen.add(target)
    return ordered


def _control_mappings_for_finding(finding: Finding) -> dict[str, list[str]]:
    title = (finding.title or "").lower()
    evidence = (finding.evidence or "").lower()
    text = f"{title} {evidence}".lower()
    if any(token in text for token in ["tls", "ssl", "cipher", "protocol", "certificate", "deprecated"]):
        return {
            "nist": ["SC-8", "SC-13", "SC-23"],
            "iso": ["A.8.24", "A.8.27", "A.8.28"],
        }
    if any(token in text for token in ["authentication", "credential", "password", "login", "bypass", "token", "session"]):
        return {
            "nist": ["AC-2", "AC-3", "IA-2"],
            "iso": ["A.5.15", "A.5.16", "A.5.17"],
        }
    if any(token in text for token in ["xss", "sql", "injection", "cross-site", "deserialization", "command"]):
        return {
            "nist": ["SI-10", "SI-11", "SC-7"],
            "iso": ["A.8.15", "A.8.28", "A.8.27"],
        }
    if any(token in text for token in ["patch", "outdated", "version", "vulnerable", "cve", "deprecated"]):
        return {
            "nist": ["SI-2", "RA-5", "CM-6"],
            "iso": ["A.8.8", "A.8.9", "A.8.19"],
        }
    if any(token in text for token in ["port", "service", "exposed", "misconfig", "default", "open"]):
        return {
            "nist": ["CM-6", "SC-7", "AC-4"],
            "iso": ["A.8.20", "A.8.21", "A.8.22"],
        }
    return {
        "nist": ["SI-2", "RA-5"],
        "iso": ["A.8.8", "A.8.9"],
    }


def _unique_controls(values: Iterable[str]) -> list[str]:
    seen: set[str] = set()
    ordered: list[str] = []
    for value in values:
        if value and value not in seen:
            seen.add(value)
            ordered.append(value)
    return ordered


def build_compliance_dashboard(findings: Iterable[Finding], *, selected_targets: list[str] | None = None) -> dict:
    filtered_findings = filter_findings(findings, selected_targets)
    grouped: dict[str, list[Finding]] = defaultdict(list)
    for finding in filtered_findings:
        target = _finding_target(finding)
        grouped[target].append(finding)

    hosts = []
    aggregated_nist: list[str] = []
    aggregated_iso: list[str] = []
    for target in sorted(grouped):
        entries = grouped[target]
        host_nist: list[str] = []
        host_iso: list[str] = []
        for entry in entries:
            mapping = _control_mappings_for_finding(entry)
            host_nist.extend(mapping.get("nist") or [])
            host_iso.extend(mapping.get("iso") or [])
        host_nist = _unique_controls(host_nist)
        host_iso = _unique_controls(host_iso)
        aggregated_nist.extend(host_nist)
        aggregated_iso.extend(host_iso)
        hosts.append(
            {
                "target": target,
                "status": "Non-compliant" if entries else "Compliant",
                "finding_count": len(entries),
                "open_findings": sum(1 for entry in entries if (entry.status or "open").lower() == "open"),
                "highest_severity": max((entry.severity or "info" for entry in entries), key=_severity_rank, default="info"),
                "controls": {"nist": host_nist, "iso": host_iso},
            }
        )

    compliant_hosts = sum(1 for host in hosts if host["status"] == "Compliant")
    non_compliant_hosts = len(hosts) - compliant_hosts
    frameworks = {
        "NIST SP 800-53 Rev. 5": {
            "controls": _unique_controls(aggregated_nist),
            "covered_hosts": sum(1 for host in hosts if host["controls"]["nist"]),
            "total_hosts": len(hosts),
        },
        "ISO/IEC 27001:2022 Annex A": {
            "controls": _unique_controls(aggregated_iso),
            "covered_hosts": sum(1 for host in hosts if host["controls"]["iso"]),
            "total_hosts": len(hosts),
        },
    }
    return {
        "summary": {
            "total_hosts": len(hosts),
            "compliant_hosts": compliant_hosts,
            "non_compliant_hosts": non_compliant_hosts,
            "open_findings": sum(host["open_findings"] for host in hosts),
        },
        "frameworks": frameworks,
        "hosts": hosts,
    }


def _styles():
    styles = getSampleStyleSheet()
    return {
        "cover_title": ParagraphStyle(
            "CoverTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=24,
            leading=30,
            textColor=colors.HexColor("#10283f"),
            alignment=TA_LEFT,
            spaceAfter=18,
        ),
        "title": ParagraphStyle(
            "ReportTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=26,
            textColor=colors.HexColor("#0d2338"),
            spaceAfter=12,
        ),
        "section": ParagraphStyle(
            "Section",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            textColor=colors.HexColor("#10283f"),
            spaceBefore=10,
            spaceAfter=8,
        ),
        "subsection": ParagraphStyle(
            "SubSection",
            parent=styles["Heading3"],
            fontName="Helvetica-Bold",
            fontSize=11,
            leading=14,
            textColor=colors.HexColor("#15324d"),
            spaceBefore=6,
            spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "Body",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9.4,
            leading=13,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#1b2f44"),
            spaceAfter=4,
        ),
        "small": ParagraphStyle(
            "Small",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=8.2,
            leading=11,
            textColor=colors.HexColor("#44576a"),
            spaceAfter=3,
        ),
        "table": ParagraphStyle(
            "TableCell",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            textColor=colors.HexColor("#1b2f44"),
        ),
        "table_header": ParagraphStyle(
            "TableHeader",
            parent=styles["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=8,
            leading=10,
            textColor=colors.white,
            alignment=TA_LEFT,
        ),
    }


def _header_drawer(company_name: str, logo_path: str | None):
    def _draw(canvas, doc):
        canvas.saveState()
        if logo_path and Path(logo_path).exists():
            try:
                canvas.drawImage(logo_path, 38, 736, width=120, height=36, preserveAspectRatio=True, mask="auto")
            except Exception:
                canvas.setFillColor(colors.HexColor("#18b4b0"))
                canvas.roundRect(38, 738, 34, 34, 8, stroke=0, fill=1)
                canvas.setFillColor(colors.white)
                canvas.setFont("Helvetica-Bold", 18)
                canvas.drawCentredString(55, 748, "V")
        else:
            canvas.setFillColor(colors.HexColor("#18b4b0"))
            canvas.roundRect(38, 738, 34, 34, 8, stroke=0, fill=1)
            canvas.setFillColor(colors.white)
            canvas.setFont("Helvetica-Bold", 18)
            canvas.drawCentredString(55, 748, "V")
        canvas.setFillColor(colors.HexColor("#10283f"))
        canvas.setFont("Helvetica-Bold", 16)
        canvas.drawString(168 if logo_path and Path(logo_path).exists() else 82, 754, company_name)
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#57697d"))
        canvas.drawString(168 if logo_path and Path(logo_path).exists() else 82, 742, "Vulnerability Assessment and Penetration Testing")
        canvas.setStrokeColor(colors.HexColor("#d7e3ec"))
        canvas.line(38, 734, 574, 734)
        canvas.setFillColor(colors.HexColor("#6c7f93"))
        canvas.setFont("Helvetica", 8)
        canvas.drawRightString(572, 24, f"Page {canvas.getPageNumber()}")
        canvas.restoreState()

    return _draw


def _table_cell(value, style):
    if isinstance(value, Paragraph):
        return value
    if value is None:
        value = ""
    paragraph = Paragraph(str(value), style)
    return paragraph


def _table(rows, col_widths):
    styles = _styles()
    normalized_rows = []
    for row_index, row in enumerate(rows):
        style = styles["table_header"] if row_index == 0 else styles["table"]
        normalized_rows.append([_table_cell(cell, style) for cell in row])
    table = Table(normalized_rows, colWidths=col_widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#10283f")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#d4dde5")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f6f9fc")]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("LEADING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("WORDWRAP", (0, 0), (-1, -1), True),
            ]
        )
    )
    return table


def _severity_table(serialized_findings: list[dict]):
    counts = Counter((item.get("severity") or "info").lower() for item in serialized_findings)
    rows = [["Severity", "Count"]]
    for severity in ("critical", "high", "medium", "low", "info"):
        rows.append([severity.title(), str(counts.get(severity, 0))])
    styles = _styles()
    normalized_rows = [[_table_cell(cell, styles["table"]) for cell in row] for row in rows]
    table = Table(normalized_rows, colWidths=[2.0 * inch, 1.2 * inch])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#10283f")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d4dde5")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f8fb")]),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("LEADING", (0, 0), (-1, -1), 11),
                ("WORDWRAP", (0, 0), (-1, -1), True),
            ]
        )
    )
    return table


def _group_by_target(serialized_findings: list[dict]) -> dict[str, list[dict]]:
    grouped: dict[str, list[dict]] = defaultdict(list)
    for item in serialized_findings:
        grouped[item.get("target") or "Unknown target"].append(item)
    return grouped


def _document_control_table(company_name: str, report_title: str, serialized_findings: list[dict], normalized_mode: str):
    targets = _scope_targets(serialized_findings)
    dates = [item.get("detected_at") for item in serialized_findings if item.get("detected_at")]
    start_date = _format_dt(min(dates)) if dates else "n/a"
    end_date = _format_dt(max(dates)) if dates else "n/a"
    rows = [
        ["Document type", report_title],
        ["Document owner", company_name],
        ["Assessment type", f"{normalized_mode.title()} security assessment"],
        ["Report date", datetime.now(timezone.utc).strftime("%d-%m-%Y")],
        ["Assessment period", f"{start_date} to {end_date}"],
        ["Targets in scope", str(len(targets))],
        ["Findings in scope", str(len(serialized_findings))],
    ]
    return _table([["Field", "Value"], *rows], [1.8 * inch, 4.6 * inch])


def _toc_table():
    rows = [
        ["1", "Introduction and assessment overview"],
        ["2", "Methodology and approach"],
        ["3", "Assessment scope"],
        ["4", "Risk level and description"],
        ["5", "Tools used during assessment"],
        ["6", "Disclaimer and assumptions"],
        ["7", "Executive summary report"],
        ["8", "Detailed report by target and finding"],
        ["9", "Compliance impact and re-test guidance"],
    ]
    return _table([["Section", "Description"], *rows], [1.0 * inch, 5.4 * inch])


def _risk_level_table():
    rows = [
        ["Critical", "CVSS 9.0-10.0", "Complete compromise or major business impact."],
        ["High", "CVSS 7.0-8.9", "Sensitive exposure or straightforward exploitation."],
        ["Medium", "CVSS 4.0-6.9", "Partial control or moderate security weakness."],
        ["Low", "CVSS 0.1-3.9", "Limited impact or hard-to-exploit weakness."],
        ["Informational", "Best practice", "Improvement item that can become security debt later."],
    ]
    return _table([["Risk", "Range", "Description"], *rows], [1.1 * inch, 1.2 * inch, 4.1 * inch])


def _tools_table(serialized_findings: list[dict]):
    source_map = {
        "openvas": "Network vulnerability scanning and service enumeration",
        "zap": "Web application crawling and vulnerability testing",
        "mobsf": "Mobile application static and metadata review",
        "network-db": "Database-backed network correlation and vulnerability enrichment",
        "platform": "Platform-side posture and workflow validation",
        "github-actions": "CI/CD or workflow security telemetry",
    }
    counts = Counter(item.get("source") or "unknown" for item in serialized_findings)
    rows = [["Tool / Engine", "Usage", "Findings"]]
    for source, count in counts.items():
        rows.append([source, source_map.get(source, "Security telemetry or enrichment source"), str(count)])
    return _table(rows, [1.5 * inch, 3.7 * inch, 1.1 * inch])


def _target_summary_table(grouped_by_asset: dict[str, list[dict]]):
    rows = [["Sl. No.", "Target", "Critical", "High", "Medium", "Low", "Info", "Total"]]
    for index, (target, entries) in enumerate(grouped_by_asset.items(), start=1):
        counts = Counter((entry.get("severity") or "info").lower() for entry in entries)
        rows.append(
            [
                str(index),
                target,
                str(counts.get("critical", 0)),
                str(counts.get("high", 0)),
                str(counts.get("medium", 0)),
                str(counts.get("low", 0)),
                str(counts.get("info", 0)),
                str(len(entries)),
            ]
        )
    return _table(rows, [0.55 * inch, 2.2 * inch, 0.65 * inch, 0.6 * inch, 0.75 * inch, 0.55 * inch, 0.55 * inch, 0.65 * inch])


def _open_port_rows(entries: list[dict]):
    seen = set()
    rows = [["Port", "Protocol", "Service", "Source"]]
    for item in entries:
        key = (item.get("port"), item.get("protocol"), item.get("service"))
        if key in seen or not item.get("port"):
            continue
        seen.add(key)
        rows.append([
            str(item.get("port") or "n/a"),
            item.get("protocol") or "tcp",
            item.get("service") or "unknown",
            item.get("source") or "scanner",
        ])
    return rows


def _finding_summary_rows(entries: list[dict]):
    rows = [["Sl. No.", "Vulnerability name", "Risk", "CVSS", "Status"]]
    for index, item in enumerate(entries, start=1):
        rows.append([
            str(index),
            item.get("title") or "Finding",
            (item.get("severity") or "info").title(),
            str(item.get("cvss_score") or "n/a"),
            item.get("status") or "open",
        ])
    return rows


def _detail_table(item: dict):
    rows = [
        ["Target", item.get("target") or "n/a"],
        ["Asset", item.get("asset_name") or "n/a"],
        ["Source", item.get("source") or "n/a"],
        ["Severity", (item.get("severity") or "info").title()],
        ["Status", item.get("status") or "open"],
        ["Verification", item.get("verification_state") or "pending"],
        ["CVE / CVSS", f"{item.get('cve_id') or 'n/a'} / {item.get('cvss_score') or 'n/a'}"],
        ["Platform", item.get("os_family") or "n/a"],
        ["CIS Benchmark", item.get("cis_benchmark") or "n/a"],
        ["Detected", item.get("detected_at") or "n/a"],
    ]
    styles = _styles()
    normalized_rows = [[_table_cell(cell, styles["table"]) for cell in row] for row in rows]
    table = Table(normalized_rows, colWidths=[1.4 * inch, 4.8 * inch])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#eff5fa")),
                ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#10283f")),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d4dde5")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("LEADING", (0, 0), (-1, -1), 10.5),
                ("WORDWRAP", (0, 0), (-1, -1), True),
            ]
        )
    )
    return table


def _reference_lines(item: dict) -> list[str]:
    refs = list(dict.fromkeys(item.get("references") or []))
    if item.get("cve_id"):
        refs.insert(0, f"https://nvd.nist.gov/vuln/detail/{item['cve_id']}")
        refs.insert(1, f"https://www.cve.org/CVERecord?id={item['cve_id']}")
    return [ref for ref in refs if ref][:6]


def report_targets(findings: Iterable[Finding]) -> list[dict]:
    grouped: dict[str, dict] = {}
    for finding in findings:
        item = _serialize_finding(finding)
        key = item.get("target") or item.get("asset_name") or str(item["id"])
        bucket = grouped.setdefault(
            key,
            {
                "target": key,
                "asset_name": item.get("asset_name") or key,
                "hostname": item.get("hostname"),
                "ip_address": item.get("ip_address"),
                "url": item.get("url"),
                "os_family": item.get("os_family"),
                "finding_count": 0,
                "highest_severity": "info",
            },
        )
        bucket["finding_count"] += 1
        if _severity_rank(item.get("severity")) > _severity_rank(bucket.get("highest_severity")):
            bucket["highest_severity"] = item.get("severity") or "info"
        for field in ("hostname", "ip_address", "url", "os_family"):
            if not bucket.get(field) and item.get(field):
                bucket[field] = item.get(field)
    return sorted(grouped.values(), key=lambda entry: (_severity_rank(entry.get("highest_severity")), entry.get("finding_count", 0), entry.get("target") or ""), reverse=True)


def filter_findings(findings: Iterable[Finding], selected_targets: list[str] | None = None) -> list[Finding]:
    selected = {value.strip().lower() for value in (selected_targets or []) if value and value.strip()}
    if not selected:
        return list(findings)
    filtered = []
    for finding in findings:
        item = _serialize_finding(finding)
        candidates = {
            (item.get("target") or "").lower(),
            (item.get("asset_name") or "").lower(),
            (item.get("hostname") or "").lower(),
            (item.get("ip_address") or "").lower(),
            (item.get("url") or "").lower(),
        }
        if candidates & selected:
            filtered.append(finding)
    return filtered


def _build_executive_summary(serialized_findings: list[dict], summary: dict) -> dict:
    severity_counts = summary.get("severity_counts", {})
    severity_ranked = sorted(
        (item for item in serialized_findings if _severity_rank(item.get("severity")) >= 3),
        key=lambda entry: _severity_rank(entry.get("severity")),
        reverse=True,
    )
    seen_titles: set[str] = set()
    priority_entries = []
    for entry in severity_ranked:
        title = entry.get("title") or "Untitled finding"
        if title in seen_titles:
            continue
        seen_titles.add(title)
        priority_entries.append(entry)
        if len(priority_entries) == 5:
            break
    top_priority_findings = [
        {
            "title": item.get("title") or "Untitled finding",
            "severity": (item.get("severity") or "info").title(),
            "target": item.get("target") or "Unspecified target",
            "cvss_score": item.get("cvss_score") or "n/a",
            "remediation": item.get("remediation") or item.get("hardening_recommendation") or "Apply the remediation guidance recorded with the finding.",
        }
        for item in priority_entries
    ]

    grouped_targets: dict[str, list[dict]] = defaultdict(list)
    for item in serialized_findings:
        grouped_targets[item.get("target") or "Unknown target"].append(item)
    highest_risk_targets = []
    for target, entries in sorted(
        grouped_targets.items(),
        key=lambda entry: (_severity_rank(max((item.get("severity") for item in entry[1]), key=_severity_rank, default="info")), len(entry[1])),
        reverse=True,
    )[:5]:
        highest_risk_targets.append(
            {
                "target": target,
                "finding_count": len(entries),
                "highest_severity": max((item.get("severity") or "info" for item in entries), key=_severity_rank, default="info"),
            }
        )

    summary_text = (
        f"The assessment identified {summary.get('total_findings', 0)} findings across the selected scope. "
        f"{summary.get('open_findings', 0)} remain open and should be tracked through a remediation plan with clear ownership."
    )
    recommendations = []
    if severity_counts.get("critical", 0):
        recommendations.append("Contain and validate all critical-risk issues immediately, especially on internet-facing or business-critical targets.")
    if severity_counts.get("high", 0):
        recommendations.append("Prioritize high-severity remediation work and confirm the fixes with re-testing before closing the items.")
    if summary.get("compliance_counts"):
        recommendations.append("Map the most significant findings to their compliance obligations and maintain evidence for audit or re-test review.")
    if not recommendations:
        recommendations.append("Review the current backlog and confirm that remediation owners and timelines have been assigned.")

    return {
        "summary_text": summary_text,
        "top_priority_findings": top_priority_findings,
        "highest_risk_targets": highest_risk_targets,
        "recommendations": recommendations,
    }


def build_report_preview(findings: Iterable[Finding], *, mode: str = "executive", selected_targets: list[str] | None = None, report_title: str | None = None) -> dict:
    filtered = filter_findings(findings, selected_targets)
    serialized = [_serialize_finding(item) for item in filtered]
    serialized.sort(
        key=lambda item: (_severity_rank(item.get("severity")), item.get("cvss_score") or 0, item.get("detected_at") or ""),
        reverse=True,
    )
    summary = summarize_findings(filtered)
    grouped_by_asset: dict[str, list[dict]] = defaultdict(list)
    for item in serialized:
        grouped_by_asset[item.get("target") or "Unknown target"].append(item)
    branding = load_report_branding()
    executive_summary = _build_executive_summary(serialized, summary)
    compliance_dashboard = build_compliance_dashboard(filtered, selected_targets=selected_targets)
    return {
        "mode": (mode or "executive").lower(),
        "report_title": report_title or f"{branding['company_name']} Security Assessment Report",
        "company_name": branding["company_name"],
        "logo_name": branding.get("logo_name"),
        "selected_targets": selected_targets or [],
        "summary": summary,
        "targets": _scope_targets(serialized),
        "findings_by_asset": [
            {
                "target": target,
                "finding_count": len(entries),
                "highest_severity": max((entry.get("severity") for entry in entries), key=_severity_rank, default="info"),
                "findings": entries[:5],
            }
            for target, entries in list(grouped_by_asset.items())[:20]
        ],
        "top_findings": serialized[:10],
        "executive_summary": executive_summary,
        "recommendations": executive_summary["recommendations"],
        "compliance_dashboard": compliance_dashboard,
    }


def _add_docx_table(document, rows: list[list[str]]):
    if not rows:
        return
    table = document.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    for row in rows:
        cells = table.add_row().cells
        for col_index, value in enumerate(row):
            paragraph = cells[col_index].paragraphs[0]
            paragraph.clear()
            paragraph.add_run().text = str(value)


def export_findings_docx(
    findings: Iterable[Finding],
    mode: str = "executive",
    *,
    selected_targets: list[str] | None = None,
    report_title: str | None = None,
    company_name: str | None = None,
) -> bytes:
    if Document is None:
        return b""

    filtered_findings = filter_findings(findings, selected_targets)
    serialized_findings = [_serialize_finding(item) for item in filtered_findings]
    serialized_findings.sort(
        key=lambda item: (_severity_rank(item.get("severity")), item.get("cvss_score") or 0, item.get("detected_at") or ""),
        reverse=True,
    )
    summary = summarize_findings(filtered_findings)
    executive_summary = _build_executive_summary(serialized_findings, summary)
    compliance_dashboard = build_compliance_dashboard(filtered_findings, selected_targets=selected_targets)
    branding = load_report_branding()
    header_name = company_name or branding["company_name"]
    title = report_title or f"{header_name} Security Assessment Report"
    normalized_mode = (mode or "executive").lower()
    mode_title = {
        "executive": "Executive Report",
        "technical": "Technical Report",
        "compliance": "Compliance Report",
    }.get(normalized_mode, "Security Report")
    mode_intro = {
        "executive": "This report provides a management-ready view of the most important risks, the affected targets, and the remediation priorities that require near-term action.",
        "technical": "This report provides detailed technical findings, validation evidence, and remediation guidance for engineering and security operations teams.",
        "compliance": "This report explains the control impact of the identified findings, the relevant benchmark references, and the evidence needed for remediation and re-test closure.",
    }.get(normalized_mode, "This report summarizes the assessed findings and recommended next steps.")

    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph("Vulnerability Assessment and Penetration Testing Report")
    document.add_paragraph(f"Report mode: {mode_title}")
    document.add_paragraph(f"Generated on {datetime.now(timezone.utc).strftime('%B %d, %Y at %H:%M UTC')}")
    document.add_paragraph(f"Company: {header_name}")
    document.add_paragraph(f"Targets in scope: {len(_scope_targets(serialized_findings))}")
    document.add_paragraph(f"Findings in scope: {summary['total_findings']}")
    document.add_paragraph(mode_intro)

    crit_cnt = summary["severity_counts"].get("critical", 0)
    high_cnt = summary["severity_counts"].get("high", 0)
    med_cnt = summary["severity_counts"].get("medium", 0)
    low_cnt = summary["severity_counts"].get("low", 0)
    targets_cnt = max(1, len(_scope_targets(serialized_findings)))
    penalty = (crit_cnt * 25 + high_cnt * 12 + med_cnt * 5 + low_cnt * 1) / (targets_cnt * 2)
    posture_score = min(100, max(20, round(100 - penalty)))

    document.add_heading("1. Executive Summary & Risk Overview", level=2)
    document.add_paragraph(executive_summary["summary_text"])
    document.add_paragraph(f"Security Posture Score: {posture_score}%")
    _add_docx_table(document, [
        ["Total Assets", "Total Findings", "Critical", "High Severity"],
        [str(targets_cnt), str(summary["total_findings"]), str(crit_cnt), str(high_cnt)]
    ])

    if normalized_mode == "executive":
        document.add_heading("2. Executive Priority Findings & Recommendations", level=2)
        for item in executive_summary["top_priority_findings"]:
            document.add_paragraph(
                f"• {item['title']} ({item['severity']}, CVSS {item['cvss_score']}) on {item['target']} — {item['remediation']}"
            )
        document.add_heading("Management Actions", level=3)
        for rec in executive_summary["recommendations"]:
            document.add_paragraph(f"• {rec}")

    elif normalized_mode == "technical":
        document.add_heading("2. Detailed Technical Findings", level=2)
        tech_rows = [["Target", "Issue Description", "Type", "CVE", "Severity", "Remediation"]]
        for item in serialized_findings:
            tech_rows.append([
                item.get("target") or "n/a",
                item.get("title") or "Finding",
                item.get("source") or "OS",
                item.get("cve_id") or "n/a",
                (item.get("severity") or "info").upper(),
                item.get("remediation") or "Apply updates.",
            ])
        _add_docx_table(document, tech_rows)

    elif normalized_mode == "compliance":
        document.add_heading("2. Regulatory Compliance Control Matrix", level=2)
        comp_rows = [["Finding Issue", "CIS", "NIST", "GDPR", "HIPAA", "SOC 2", "ISO 27001"]]
        for item in serialized_findings:
            issue = item.get("title") or ""
            sev = (item.get("severity") or "").upper()
            comp_rows.append([
                issue,
                "X",
                "X" if sev in {"CRITICAL", "HIGH"} else "-",
                "X" if "Password" in issue or "Exposed" in issue else "-",
                "X" if sev == "CRITICAL" else "-",
                "X" if "CORS" in issue or "Credentials" in issue else "-",
                "X"
            ])
        _add_docx_table(document, comp_rows)

    document.add_heading("1. Introduction", level=2)
    document.add_paragraph(
        "This report summarizes the vulnerability assessment and penetration testing results for the selected in-scope assets."
    )
    document.add_paragraph(
        "The objective of the assessment is to identify exploitable weaknesses, explain their business and technical impact, and provide clear remediation guidance that can be validated during re-test."
    )

    document.add_heading("2. Methodology and approach", level=2)
    document.add_paragraph(
        "The assessment combines network, web, mobile, and platform-originated telemetry where available. Findings are normalized, correlated, and severity-ranked using source evidence, CVSS context, and platform enrichment."
    )
    document.add_paragraph(
        "The workflow includes discovery, fingerprinting, vulnerability validation, evidence capture, enrichment, reporting, and re-test preparation."
    )

    document.add_heading("3. Assessment scope", level=2)
    targets = _scope_targets(serialized_findings)
    if not targets:
        document.add_paragraph("No targets were selected for this report scope.")
    else:
        _add_docx_table(document, [["Sl. No.", "Target", "OS / Platform"]] + [[str(index), target, (next((item.get("os_family") or "n/a" for item in serialized_findings if item.get("target") == target), "n/a"))] for index, target in enumerate(targets, start=1)])

    document.add_heading("4. Risk level and description", level=2)
    _add_docx_table(
        document,
        [
            ["Risk", "Range", "Description"],
            ["Critical", "CVSS 9.0-10.0", "Complete compromise or major business impact."],
            ["High", "CVSS 7.0-8.9", "Sensitive exposure or straightforward exploitation."],
            ["Medium", "CVSS 4.0-6.9", "Partial control or moderate security weakness."],
            ["Low", "CVSS 0.1-3.9", "Limited impact or hard-to-exploit weakness."],
            ["Informational", "Best practice", "Improvement item that can become security debt later."],
        ],
    )

    document.add_heading("5. Tools used during assessment", level=2)
    source_map = {
        "openvas": "Network vulnerability scanning and service enumeration",
        "zap": "Web application crawling and vulnerability testing",
        "mobsf": "Mobile application static and metadata review",
        "network-db": "Database-backed network correlation and vulnerability enrichment",
        "platform": "Platform-side posture and workflow validation",
        "github-actions": "CI/CD or workflow security telemetry",
    }
    counts = Counter(item.get("source") or "unknown" for item in serialized_findings)
    rows = [["Tool / Engine", "Usage", "Findings"]]
    for source, count in counts.items():
        rows.append([source, source_map.get(source, "Security telemetry or enrichment source"), str(count)])
    _add_docx_table(document, rows)

    document.add_heading("6. Disclaimer and assumptions", level=2)
    document.add_paragraph(
        "This report reflects the state of the assessed assets at the time of testing and within the limits of the configured scan depth, selected validation routines, available credentials, and reachable services."
    )
    document.add_paragraph(
        "Before implementing remediation, confirm that the affected service is required for business use, prepare backup and rollback plans, and schedule any disruptive changes through the appropriate operational process."
    )

    document.add_heading("7. Executive summary report", level=2)
    document.add_paragraph(executive_summary["summary_text"])
    document.add_paragraph(
        f"A total of {summary['total_findings']} findings are in scope. {summary['open_findings']} remain open."
    )
    document.add_paragraph("Priority findings", style="Intense Quote")
    for item in executive_summary["top_priority_findings"]:
        document.add_paragraph(
            f"• {item['title']} ({item['severity']}, CVSS {item['cvss_score']}) on {item['target']} — {item['remediation']}"
        )
    document.add_paragraph("Management actions", style="Intense Quote")
    for recommendation in executive_summary["recommendations"]:
        document.add_paragraph(f"• {recommendation}")
    _add_docx_table(document, [["Severity", "Count"], ["Critical", str(summary["severity_counts"].get("critical", 0))], ["High", str(summary["severity_counts"].get("high", 0))], ["Medium", str(summary["severity_counts"].get("medium", 0))], ["Low", str(summary["severity_counts"].get("low", 0))], ["Info", str(summary["severity_counts"].get("info", 0))]])

    document.add_heading("8. Detailed report", level=2)
    grouped_by_target = defaultdict(list)
    for item in serialized_findings:
        grouped_by_target[item.get("target") or "Unknown target"].append(item)
    for target, entries in grouped_by_target.items():
        document.add_heading(f"Target: {target}", level=3)
        document.add_paragraph("Summary of findings")
        _add_docx_table(document, [["Sl. No.", "Vulnerability name", "Risk", "CVSS", "Status"]] + [[str(index), entry.get("title") or "Finding", (entry.get("severity") or "info").title(), str(entry.get("cvss_score") or "n/a"), entry.get("status") or "open"] for index, entry in enumerate(entries, start=1)])
        for item in entries[:3]:
            document.add_paragraph(f"• {item.get('title')} — {item.get('severity') or 'info'}")

    document.add_heading("8.1 Compliance posture", level=3)
    if compliance_dashboard["hosts"]:
        _add_docx_table(
            document,
            [["Host", "Status", "Open findings", "NIST controls", "ISO controls"]]
            + [
                [
                    host["target"],
                    host["status"],
                    str(host["open_findings"]),
                    ", ".join(host["controls"]["nist"]),
                    ", ".join(host["controls"]["iso"]),
                ]
                for host in compliance_dashboard["hosts"]
            ],
        )
    else:
        document.add_paragraph("No host-level compliance posture information was available for the selected scope.")

    document.add_heading("9. Compliance impact and re-test guidance", level=2)
    compliance_counts = Counter(tag for item in serialized_findings for tag in (item.get("compliance_map") or []))
    if compliance_counts:
        for tag, count in compliance_counts.most_common(10):
            document.add_paragraph(f"• {tag}: {count} related finding(s)")
    else:
        document.add_paragraph("No compliance mappings were available for the current report scope.")
    document.add_paragraph("Re-test guidance")
    for line in [
        "Re-test each affected target after remediation is completed.",
        "Validate that the vulnerability is no longer reproducible and that the service matches the intended secure state.",
        "Capture screenshots, configuration outputs, patch evidence, or scanner proof for closure.",
        "Where a fix cannot be implemented immediately, document the compensating control, risk owner, and next review date.",
    ]:
        document.add_paragraph(f"• {line}")

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def export_findings_pdf(
    findings: Iterable[Finding],
    mode: str = "executive",
    *,
    selected_targets: list[str] | None = None,
    report_title: str | None = None,
    company_name: str | None = None,
    author_name: str | None = None,
    logo_path: str | None = None,
) -> bytes:
    if colors is None:
        return b"%PDF-1.4\n% Report engine unavailable\n"

    filtered_findings = filter_findings(findings, selected_targets)
    serialized_findings = [_serialize_finding(item) for item in filtered_findings]
    serialized_findings.sort(
        key=lambda item: (_severity_rank(item.get("severity")), item.get("cvss_score") or 0, item.get("detected_at") or ""),
        reverse=True,
    )
    summary = summarize_findings(filtered_findings)
    executive_summary = _build_executive_summary(serialized_findings, summary)
    styles = _styles()
    buffer = io.BytesIO()
    branding = load_report_branding()
    header_name = company_name or branding["company_name"]
    normalized_mode = (mode or "executive").lower()
    title_name = report_title or f"Vulnerability Assessment {normalized_mode.title()} Report"
    author_str = author_name or "Lead Security Auditor"

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=36,
        rightMargin=36,
        topMargin=40,
        bottomMargin=36,
        title=title_name,
        author=author_str,
    )

    crit_cnt = summary["severity_counts"].get("critical", 0)
    high_cnt = summary["severity_counts"].get("high", 0)
    med_cnt = summary["severity_counts"].get("medium", 0)
    low_cnt = summary["severity_counts"].get("low", 0)
    targets_cnt = max(1, len(_scope_targets(serialized_findings)))
    penalty = (crit_cnt * 25 + high_cnt * 12 + med_cnt * 5 + low_cnt * 1) / (targets_cnt * 2)
    posture_score = min(100, max(20, round(100 - penalty)))

    story = []

    # 1. Document Header (Matches Live Preview Paper View Top Banner)
    story.append(Paragraph(header_name, ParagraphStyle("CompHeader", fontName="Helvetica-Bold", fontSize=22, leading=26, textColor=colors.HexColor("#0f172a"))))
    story.append(Paragraph(title_name, ParagraphStyle("ReportSubTitle", fontName="Helvetica-Bold", fontSize=14, leading=18, textColor=colors.HexColor("#2563eb"), spaceAfter=8)))
    story.append(Paragraph(f"<b>Author:</b> {author_str} &nbsp;&nbsp;|&nbsp;&nbsp; <b>Generated:</b> {datetime.now(timezone.utc).strftime('%d/%m/%Y')} &nbsp;&nbsp;|&nbsp;&nbsp; <b>Mode:</b> {normalized_mode.upper()}", styles["small"]))
    story.append(Spacer(1, 14))

    # 2. Section 1: Executive Summary & Risk Overview
    story.append(Paragraph("1. Executive Summary & Risk Overview", styles["section"]))
    story.append(Paragraph(executive_summary["summary_text"], styles["body"]))
    story.append(Spacer(1, 8))

    # 4 Metric Cards Table (Matches Live Preview Grid Cards)
    metric_table_data = [
        ["TOTAL ASSETS", "TOTAL FINDINGS", "CRITICAL", "HIGH SEVERITY"],
        [str(targets_cnt), str(summary["total_findings"]), str(crit_cnt), str(high_cnt)],
    ]
    metric_table = Table(metric_table_data, colWidths=[1.75 * inch, 1.75 * inch, 1.75 * inch, 1.75 * inch])
    metric_table.setStyle(
        TableStyle([
            ("BACKGROUND", (0, 0), (1, -1), colors.HexColor("#f8fafc")),
            ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#fef2f2")),
            ("BACKGROUND", (3, 0), (3, -1), colors.HexColor("#fffbe0")),
            ("TEXTCOLOR", (0, 0), (1, 0), colors.HexColor("#64748b")),
            ("TEXTCOLOR", (2, 0), (2, 0), colors.HexColor("#dc2626")),
            ("TEXTCOLOR", (3, 0), (3, 0), colors.HexColor("#d97706")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 8),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("FONTNAME", (0, 1), (-1, 1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 1), (-1, 1), 16),
            ("TEXTCOLOR", (0, 1), (1, 1), colors.HexColor("#0f172a")),
            ("TEXTCOLOR", (2, 1), (2, 1), colors.HexColor("#991b1b")),
            ("TEXTCOLOR", (3, 1), (3, 1), colors.HexColor("#92400e")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
            ("TOPPADDING", (0, 0), (-1, -1), 8),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ])
    )
    story.append(metric_table)
    story.append(Spacer(1, 10))

    # Security Posture Gauge Score Box
    score_box_data = [
        [f"Security Posture Score: {posture_score}%\nEvaluated across active infrastructure, web applications, and network endpoints."]
    ]
    score_table = Table(score_box_data, colWidths=[7.0 * inch])
    score_table.setStyle(
        TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f8fafc")),
            ("BORDER", (0, 0), (-1, -1), 1, colors.HexColor("#cbd5e1")),
            ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#0f172a")),
            ("FONTNAME", (0, 0), (-1, -1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
            ("TOPPADDING", (0, 0), (-1, -1), 10),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
            ("LEFTPADDING", (0, 0), (-1, -1), 14),
        ])
    )
    story.append(score_table)
    story.append(Spacer(1, 16))

    # Mode-specific sections:
    if normalized_mode == "technical" or normalized_mode == "executive":
        story.append(Paragraph(f"2. Detailed Technical Findings ({len(serialized_findings)})", styles["section"]))
        if serialized_findings:
            tech_rows = [["Target", "Issue Description", "Type", "CVE", "Severity", "Remediation"]]
            for item in serialized_findings:
                tech_rows.append([
                    item.get("target") or item.get("asset_name") or "n/a",
                    item.get("title") or "Finding",
                    item.get("source") or "OS",
                    item.get("cve_id") or "n/a",
                    (item.get("severity") or "info").upper(),
                    item.get("remediation") or "Apply updates.",
                ])
            tech_table = _table(tech_rows, [1.3 * inch, 2.0 * inch, 0.7 * inch, 0.9 * inch, 0.8 * inch, 1.3 * inch])
            story.append(tech_table)
        else:
            story.append(Paragraph("No findings discovered for this report.", styles["body"]))
        story.append(Spacer(1, 16))

    if normalized_mode == "compliance" or normalized_mode == "executive":
        story.append(Paragraph("3. Regulatory Compliance Control Matrix", styles["section"]))
        if serialized_findings:
            comp_rows = [["Finding Issue", "CIS", "NIST", "GDPR", "HIPAA", "SOC 2", "ISO 27001"]]
            for item in serialized_findings:
                issue = item.get("title") or ""
                sev = (item.get("severity") or "").upper()
                comp_rows.append([
                    issue,
                    "X",
                    "X" if sev in {"CRITICAL", "HIGH"} else "-",
                    "X" if "Password" in issue or "Exposed" in issue else "-",
                    "X" if sev == "CRITICAL" else "-",
                    "X" if "CORS" in issue or "Credentials" in issue else "-",
                    "X"
                ])
            comp_table = _table(comp_rows, [2.2 * inch, 0.8 * inch, 0.8 * inch, 0.8 * inch, 0.8 * inch, 0.8 * inch, 0.8 * inch])
            story.append(comp_table)
        else:
            story.append(Paragraph("No compliance violations identified in the current scope.", styles["body"]))

    doc.build(story)
    return buffer.getvalue()
